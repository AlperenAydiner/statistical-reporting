"""python-docx post-processing pass, applied after pandoc has assembled the
document. This is the step the plan calls non-optional: pandoc's docx
writer cannot emit merged cells and gives every table a full grid of
borders, neither of which is APA 7 table style (no vertical rules; a rule
above the header, one under the header row, one under the table; numeric
columns aligned; the header row repeating across page breaks).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from statrep.analysis.results import TableSpec

_NONE_BORDER = {"sz": "0", "val": "nil", "color": "auto"}
_RULE_BORDER = {"sz": "6", "val": "single", "color": "000000"}  # 0.75pt

# CT_TcBorders requires its child elements in this exact sequence — the XSD
# schema rejects an otherwise-valid <w:tcBorders> if e.g. <w:left> is
# written before <w:top>. Caught by validating real output against the
# ISO-29500 schema; a plain dict-order append (whatever order the caller's
# kwargs happened to be in) fails silently in Word but fails loudly there.
_BORDER_EDGE_ORDER = ("top", "left", "bottom", "right", "insideH", "insideV", "tl2br", "tr2bl")


def _set_cell_borders(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders_el = tc_pr.find(qn("w:tcBorders"))
    if borders_el is None:
        borders_el = OxmlElement("w:tcBorders")
        tc_pr.append(borders_el)
    for edge in _BORDER_EDGE_ORDER:
        if edge not in edges:
            continue
        spec = edges[edge]
        tag = f"w:{edge}"
        el = borders_el.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders_el.append(el)
        for key, val in spec.items():
            el.set(qn(f"w:{key}"), val)


def _mark_header_row_repeating(table) -> None:
    try:
        table.rows[0].header_row = True
    except AttributeError:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:tblHeader"))


def apply_apa_style(table, numeric_columns: list[int]) -> None:
    """No vertical rules anywhere; horizontal rules above the header, under
    the header row, and under the table's last row. Numeric-column data
    cells right-aligned. Header row repeats across page breaks."""
    n_rows = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            edges = {"left": _NONE_BORDER, "right": _NONE_BORDER}
            edges["top"] = _RULE_BORDER if r_idx == 0 else _NONE_BORDER
            edges["bottom"] = _RULE_BORDER if r_idx in (0, n_rows - 1) else _NONE_BORDER
            _set_cell_borders(cell, **edges)
            if c_idx in numeric_columns and r_idx > 0:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _mark_header_row_repeating(table)


def apply_merged_header(table, spans: list[tuple[int, int, str]]) -> None:
    """Merge a run of header cells into one labeled span — the one thing
    pandoc's docx writer cannot produce at all (APA hierarchical
    regression's grouped columns, SPSS's nested pivot headers)."""
    if not spans:
        return
    header_row = table.rows[0]
    for col_start, col_end, label in spans:
        merged = header_row.cells[col_start]
        for c in range(col_start + 1, col_end + 1):
            merged = merged.merge(header_row.cells[c])
        merged.text = label


def postprocess(docx_path: str | Path, table_specs: list[TableSpec]) -> Path:
    """Match the docx's tables to ``table_specs`` by document order (pandoc
    inserts tables in the same order they appear in the source Markdown)
    and apply APA styling to each."""
    docx_path = Path(docx_path)
    doc = Document(docx_path)
    for docx_table, spec in zip(doc.tables, table_specs):
        apply_apa_style(docx_table, spec.numeric_columns)
        apply_merged_header(docx_table, spec.merged_header_spans)
    doc.save(docx_path)
    return docx_path
